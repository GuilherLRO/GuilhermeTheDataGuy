"""
Generate DOCX documentation files from Markdown sources.

This script converts the markdown documentation files in the Nike documentation
folder into professionally formatted Word documents.

Requirements:
    pip install python-docx markdown

Usage:
    python generate_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def setup_document_styles(doc: Document) -> None:
    """Configure document styles for professional formatting."""
    
    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 82, 147)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(4)
    
    # Heading 3
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(68, 68, 68)
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after = Pt(4)
    
    # Normal text
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)


def add_table_from_markdown(doc: Document, table_lines: list[str]) -> None:
    """Parse markdown table and add to document."""
    
    # Parse table rows
    rows = []
    for line in table_lines:
        line = line.strip()
        if line.startswith('|') and not re.match(r'^\|[\s\-:|]+\|$', line):
            # Split by | and clean up
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Fill table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                # Clean markdown formatting
                cell_text = re.sub(r'`([^`]+)`', r'\1', cell_text)  # Remove backticks
                cell.text = cell_text
                
                # Style header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        paragraph.runs[0].bold = True if paragraph.runs else False
                        # Set header background color
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8"/>')
                        cell._tc.get_or_add_tcPr().append(shading)
                
                # Set font
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
    
    # Add spacing after table
    doc.add_paragraph()


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    """Add a formatted code block to the document."""
    code_text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    # Light gray background effect via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._p.get_or_add_pPr().append(shading)


def process_inline_formatting(paragraph, text: str) -> None:
    """Process inline markdown formatting (bold, italic, code)."""
    # Pattern to match **bold**, *italic*, `code`
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def convert_markdown_to_docx(md_path: Path, output_path: Path) -> None:
    """Convert a markdown file to a formatted DOCX document."""
    
    print(f"Converting: {md_path.name} -> {output_path.name}")
    
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    setup_document_styles(doc)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    
    # Use first H1 as title
    first_h1 = True
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Horizontal rule (---)
        if re.match(r'^-{3,}$', stripped):
            # Add a subtle separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Headers
        if stripped.startswith('#'):
            header_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if header_match:
                level = len(header_match.group(1))
                header_text = header_match.group(2)
                
                if level == 1:
                    if first_h1:
                        doc.add_heading(header_text, 0)  # Title
                        first_h1 = False
                    else:
                        doc.add_heading(header_text, 1)
                elif level == 2:
                    doc.add_heading(header_text, 2)
                elif level == 3:
                    doc.add_heading(header_text, 3)
                else:
                    doc.add_heading(header_text, min(level, 3))
                
                i += 1
                continue
        
        # Code blocks
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # Skip closing ```
            continue
        
        # Tables
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_markdown(doc, table_lines)
            continue
        
        # List items
        if re.match(r'^[-*]\s+', stripped):
            list_text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, list_text)
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s+', stripped):
            list_text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, list_text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        process_inline_formatting(p, stripped)
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"  ✓ Saved: {output_path}")


def main():
    """Main function to convert all documentation files."""
    
    # Get script directory
    script_dir = Path(__file__).parent
    
    # Define files to convert
    md_files = [
        'FoP Definitions.md',
        'L1-L3 Hierarchy Definitions.md',
        'Gender Classification Definitions.md',
        'Values Reference Tables.md',
    ]
    
    print("=" * 60)
    print("📄 Documentation DOCX Generator")
    print("=" * 60)
    print()
    
    # Create output directory for docx files
    output_dir = script_dir / 'docx'
    output_dir.mkdir(exist_ok=True)
    print(f"Output directory: {output_dir}")
    print()
    
    # Convert each file
    converted = 0
    for md_file in md_files:
        md_path = script_dir / md_file
        if md_path.exists():
            output_name = md_path.stem + '.docx'
            output_path = output_dir / output_name
            try:
                convert_markdown_to_docx(md_path, output_path)
                converted += 1
            except Exception as e:
                print(f"  ✗ Error converting {md_file}: {e}")
        else:
            print(f"  ⚠ File not found: {md_file}")
    
    print()
    print("=" * 60)
    print(f"✅ Converted {converted}/{len(md_files)} files")
    print(f"📁 Output: {output_dir}")
    print("=" * 60)


if __name__ == '__main__':
    main()

